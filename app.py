import os
import streamlit as st
from dotenv import load_dotenv
from langchain.schema import SystemMessage, HumanMessage
from langchain_groq import ChatGroq
import replicate
from fpdf import FPDF
from docx import Document
import requests
from docx.shared import Inches

import pickle
from sklearn.feature_extraction.text import TfidfVectorizer
from sklearn.naive_bayes import MultinomialNB
from sklearn.pipeline import make_pipeline
import nltk
nltk.download('punkt')



# Load API key
load_dotenv()
api_key = os.getenv("GROQ_API_KEY")
assert api_key is not None, "Set your Groq API key in .env"
REPLICATE_API_TOKEN = os.getenv("REPLICATE_API_TOKEN")
assert REPLICATE_API_TOKEN is not None, "Set REPLICATE_API_TOKEN in .env"

# Initialize Groq LLM (use Mixtral or LLaMA3 models)
llm = ChatGroq(
    groq_api_key=api_key,
    model_name="llama3-70b-8192",  # or "llama3-70b-8192"
    temperature=0.8
)

# Function to generate an image using Replicate
@st.cache_data(show_spinner=False)
def generate_image(prompt, num_images=1):
    os.environ["REPLICATE_API_TOKEN"] = REPLICATE_API_TOKEN
    images = []
    for _ in range(num_images):
        output = replicate.run(
            "black-forest-labs/flux-pro",
            input={
                "prompt": prompt,
                "guidance_scale": 7.5,
                "num_inference_steps": 25
            }
        )
        images.append(str(output))  # Ensure it’s serializable
    return images


def clean_text(text):
    replacements = {
        "–": "-",  # en-dash to hyphen
        "—": "-",  # em-dash to hyphen
        "“": '"',  # left double quote
        "”": '"',  # right double quote
        "‘": "'",  # left single quote
        "’": "'",  # right single quote
    }
    for old, new in replacements.items():
        text = text.replace(old, new)
    return text



# Save blog as PDF
def save_as_pdf(text, filename="blog.pdf"):
    pdf = FPDF()
    pdf.add_page()
    pdf.set_auto_page_break(auto=True, margin=15)
    pdf.set_font("Arial", size=12)
    clean = clean_text(text)  # 👈 Clean text

    #Add Text
    for line in clean.split('\n'):
        pdf.multi_cell(0, 10, line)

     # Add images (optional)
    if image_urls:
        for idx, url in enumerate(image_urls, start=1):
            response = requests.get(url)
            image_path = f"temp_image_{idx}.jpg"
            with open(image_path, "wb") as f:
                f.write(response.content)
            pdf.add_page()
            pdf.image(image_path, x=10, y=20, w=180)  # Adjust position/size as needed

    pdf.output(filename)
    return filename


# Save blog as Word Document
def save_as_docx(text, filename="blog.docx"):
    doc = Document()
    #Add text
    for line in text.split('\n'):
        doc.add_paragraph(line)

    # Add images (optional)
    if image_urls:
        for idx, url in enumerate(image_urls, start=1):
            response = requests.get(url)
            image_path = f"temp_image_{idx}.jpg"
            with open(image_path, "wb") as f:
                f.write(response.content)
            doc.add_picture(image_path, width=Inches(5))  # Adjust width as needed
    doc.save(filename)
    return filename


# Streamlit UI
st.set_page_config(layout="wide")
st.title("BlogCraft ✍️🤖")
st.subheader("Generate AI-powered blogs with images")

@st.cache_resource
def get_text_classifier():
    import pickle
    with open("blog_category_model.pkl", "rb") as f:
        model = pickle.load(f)
    categories = model.classes_  # list of category names
    return model, categories


with st.sidebar:
    st.header("Inputs")
    with st.form("input_form"):
        title = st.text_input("Blog Title", "")
        keywords = st.text_area("Keywords (comma-separated)", "")
        word_count = st.slider("Approx Length", 250, 2000, 500, 250)
        num_images = st.slider("Number of Images", 1, 5, 1)
        generate = st.form_submit_button("Generate Blog")
        if title and keywords:
            classifier, category_names = get_text_classifier()
            input_text = f"{title} {keywords}"
            predicted = classifier.predict([input_text])[0]
            predicted_category = predicted[0:20]
            st.markdown(f"📂 **Predicted Blog Category**: `{predicted_category}`")

        reset = st.form_submit_button("Reset")

if reset:
    st.rerun()

if generate and title:
    prompt = (
        f"Please write a detailed, engaging blog post titled '{title}'. "
        f"Include these keywords: {keywords}. "
        f"Aim for about {word_count} words. "
        "Use a conversational, professional tone and structure it with headings."
    )

    messages = [
        SystemMessage(content="You are a professional blog writer AI."),
        HumanMessage(content=prompt),
    ]

    with st.spinner("Generating blog content..."):
        response = llm.invoke(messages)
        blog_content = response.content

    st.subheader("📝 Generated Blog")
    st.write(blog_content)

    with st.spinner("Generating images..."):
        image_urls = generate_image(title, num_images=num_images)
        for idx, url in enumerate(image_urls, start=1):
            st.image(url, caption=f"AI-generated Image #{idx}")


    # Save options
    pdf_path = save_as_pdf(blog_content)
    docx_path = save_as_docx(blog_content)

    with open(pdf_path, "rb") as f:
        st.download_button("Download PDF", f, file_name="blog.pdf")

    with open(docx_path, "rb") as f:
        st.download_button("Download Word", f, file_name="blog.docx")

elif generate:
    st.error("Please enter a blog title.")
